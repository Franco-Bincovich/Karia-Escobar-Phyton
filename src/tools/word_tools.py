"""
Generación de documentos Word (.docx) con formatos formales.
Tipos: general, oficio, circular, acta, respuesta.
"""

import os

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.utils.app_error import AppError
from src.utils.logger import get_logger

logger = get_logger("wordTools")

TMP_DIR = "/tmp"


def _add_titulo(doc: Document, texto: str):
    """Agrega un título centrado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(14)


def _add_parrafo(doc: Document, texto: str, bold: bool = False):
    """Agrega un párrafo con estilo básico."""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = bold
    run.font.size = Pt(11)


def _add_label(doc: Document, label: str, valor: str):
    """Agrega un par label: valor en negrita/normal."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_valor = p.add_run(valor)
    run_valor.font.size = Pt(11)


def _build_general(doc: Document, titulo: str, contenido: str, meta: dict):
    """Builder para documento general."""
    _add_titulo(doc, titulo)
    for linea in contenido.split("\n"):
        if linea.strip():
            _add_parrafo(doc, linea.strip())


def _build_oficio(doc: Document, titulo: str, contenido: str, meta: dict):
    """Builder para oficio formal."""
    _add_titulo(doc, meta.get("organismo", "MUNICIPALIDAD DE ESCOBAR"))
    _add_label(doc, "OFICIO N°", meta.get("numero", "___/____"))
    _add_label(doc, "Destinatario", meta.get("destinatario", ""))
    _add_label(doc, "Ref", titulo)
    _add_label(doc, "Fecha", meta.get("fecha", ""))
    doc.add_paragraph()
    for linea in contenido.split("\n"):
        if linea.strip():
            _add_parrafo(doc, linea.strip())
    doc.add_paragraph()
    _add_parrafo(doc, meta.get("firmante", ""), bold=True)


def _build_circular(doc: Document, titulo: str, contenido: str, meta: dict):
    """Builder para circular."""
    _add_titulo(doc, meta.get("organismo", "MUNICIPALIDAD DE ESCOBAR"))
    _add_label(doc, "CIRCULAR N°", meta.get("numero", "___/____"))
    _add_label(doc, "Destinatario", meta.get("destinatario", "TODO EL PERSONAL"))
    _add_label(doc, "Ref", titulo)
    _add_label(doc, "Fecha", meta.get("fecha", ""))
    doc.add_paragraph()
    for i, linea in enumerate(contenido.split("\n"), 1):
        if linea.strip():
            _add_parrafo(doc, f"{i}. {linea.strip()}")


def _build_acta(doc: Document, titulo: str, contenido: str, meta: dict):
    """Builder para acta de reunión."""
    _add_titulo(doc, "ACTA DE REUNIÓN")
    _add_label(doc, "Fecha", meta.get("fecha", ""))
    _add_label(doc, "Lugar", meta.get("lugar", ""))
    _add_label(doc, "Presentes", meta.get("presentes", ""))
    doc.add_paragraph()
    for linea in contenido.split("\n"):
        if linea.strip():
            _add_parrafo(doc, linea.strip())


def _build_respuesta(doc: Document, titulo: str, contenido: str, meta: dict):
    """Builder para nota de respuesta."""
    _add_titulo(doc, meta.get("organismo", "MUNICIPALIDAD DE ESCOBAR"))
    _add_titulo(doc, "NOTA DE RESPUESTA")
    _add_label(doc, "Ref", meta.get("referencia", titulo))
    _add_label(doc, "Expediente", meta.get("expediente", ""))
    _add_label(doc, "Fecha", meta.get("fecha", ""))
    doc.add_paragraph()
    for linea in contenido.split("\n"):
        if linea.strip():
            _add_parrafo(doc, linea.strip())
    doc.add_paragraph()
    _add_parrafo(doc, meta.get("firmante", ""), bold=True)


_BUILDERS = {
    "general": _build_general,
    "oficio": _build_oficio,
    "circular": _build_circular,
    "acta": _build_acta,
    "respuesta": _build_respuesta,
}


async def generar_word(
    nombreArchivo: str, titulo: str, contenido: str, user_id: str,
    tipoDocumento: str = "general", metadatos: dict = None, **_,
) -> str:
    """
    Genera un documento Word (.docx) con formato según el tipo.

    Returns:
        Ruta del archivo generado en /tmp.
    """
    try:
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1)

        builder = _BUILDERS.get(tipoDocumento, _BUILDERS["general"])
        builder(doc, titulo, contenido, metadatos or {})

        ruta = os.path.join(TMP_DIR, f"{user_id}_{nombreArchivo}.docx")
        doc.save(ruta)
        logger.info("Word generado ruta=%s tipo=%s", ruta, tipoDocumento)
        return ruta
    except AppError:
        raise
    except Exception as err:
        logger.error("Error al generar Word: %s", str(err))
        raise AppError(f"Error al generar Word: {err}", "EXPORT_ERROR", 500)
